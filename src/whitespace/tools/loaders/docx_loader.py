"""DOCX loader (python-docx)."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)


class DocxLoader:
    """Extracts paragraph and table text from a DOCX file."""

    def load(self, path: str) -> str:
        if not Path(path).is_file():
            raise ValueError(f"DOCX not found: {path}")
        parts: list[str] = []
        try:
            document = Document(path)
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts.append(text)
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    line = " | ".join(c for c in cells if c)
                    if line:
                        parts.append(line)
        except Exception as exc:
            logger.warning("DocxLoader: best-effort extraction for %s (%s)", path, exc)
        return "\n".join(parts)
